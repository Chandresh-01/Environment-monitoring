import os
import numpy as np
import matplotlib.pyplot as plt
from tensorflow.keras.preprocessing import image
from collections import Counter, defaultdict
from tensorflow.keras.models import load_model
from sklearn.linear_model import LinearRegression
from docx import Document
from docx.shared import Inches

# Load models
env_model = load_model("models/environment_cnn_model.keras")
haze_model = load_model("models/haze_cnn_model.keras")

# Class mappings
env_class_labels = {
    'airport': 0, 'bareland': 1, 'beach': 2, 'bridge': 3, 'center': 4, 'church': 5,
    'commercial': 6, 'desert': 7, 'forest': 8, 'golf_course': 9, 'industrial_area': 10,
    'lake': 11, 'residential': 12, 'parking_lot': 13, 'river': 14, 'snowberg': 15,
    'glacier': 16, 'airplane': 17, 'runway': 18, 'stadium': 19
}
env_labels = dict((v, k) for k, v in env_class_labels.items())

env_map = {
    'forest': 'Green Cover', 'golf_course': 'Green Cover',
    'river': 'Waterbody', 'lake': 'Waterbody', 'beach': 'Waterbody',
    'bridge': 'Urban Area', 'industrial_area': 'Pollution Zone',
    'residential': 'Urban Area', 'parking_lot': 'Urban Area',
    'bareland': 'Arid Zone', 'desert': 'Arid Zone',
    'glacier': 'Climate Impacted', 'snowberg': 'Climate Impacted',
    'center': 'Urban Area', 'airport': 'Urban Area', 'airplane': 'Other',
    'commercial': 'Urban Area', 'church': 'Urban Area', 'runway': 'Infrastructure',
    'stadium': 'Urban Area'
}

haze_labels = {0: 'clear', 1: 'hazy'}

def classify_env_type(label):
    return env_map.get(label, 'Other')

def analyze_uploaded_images(upload_folder):
    env_counts = defaultdict(Counter)
    haze_counts = defaultdict(Counter)

    for year in os.listdir(upload_folder):
        year_path = os.path.join(upload_folder, year)
        if not os.path.isdir(year_path): continue

        if year.isdigit() and len(year) == 4:
            final_year = year
        else:
            digits = ''.join(filter(str.isdigit, year))
            final_year = digits if len(digits) == 4 else None
            if not final_year:
                print(f"Skipping invalid year folder: {year}")
                continue

        final_year = str(final_year)  # Ensure string keys

        for img_name in os.listdir(year_path):
            print(f"Found image: {img_name} in year: {final_year}")

            if not img_name.lower().endswith(('jpg', 'jpeg', 'png')): continue
            try:
                img_path = os.path.join(year_path, img_name)
                img = image.load_img(img_path, target_size=(150, 150))
                img_array = image.img_to_array(img) / 255.0
                img_array = np.expand_dims(img_array, axis=0)

                env_pred = env_model.predict(img_array, verbose=0)
                env_class = np.argmax(env_pred)
                env_label = env_labels.get(env_class, 'unknown')
                env_type = classify_env_type(env_label)
                env_counts[final_year][env_type] += 1

                haze_pred = haze_model.predict(img_array, verbose=0)

                if haze_pred.shape[-1] == 1:
                # Binary classification (sigmoid output)
                    haze_class = 1 if haze_pred[0][0] >= 0.5 else 0
                else:
                    # Categorical classification (softmax output)
                    haze_class = np.argmax(haze_pred)

                haze_label = haze_labels[haze_class]

                haze_counts[final_year][haze_label] += 1

                print(f"Year: {final_year}, Env: {env_type}, Haze: {haze_label}")

            except Exception as e:
                print(f"Error processing {img_name}: {e}")

    return env_counts, haze_counts

def generate_plot(env_counts, haze_counts, output_path):
    years_str = sorted(env_counts.keys(), key=int)
    years = list(map(int, years_str))
    categories = sorted({c for counts in env_counts.values() for c in counts})

    for cat in categories:
        values = [env_counts[y].get(cat, 0) for y in years_str]
        plt.plot(years, values, marker='o', label=cat)

    for haze_type in ['clear', 'hazy']:
        haze_values = [haze_counts[y].get(haze_type, 0) for y in years_str]
        plt.plot(years, haze_values, linestyle='--', marker='x', label=f"Haze: {haze_type}")

    plt.title("Environmental & Haze Trends Over Years")
    plt.xlabel("Year")
    plt.ylabel("Image Count")
    plt.legend()
    plt.grid(True)
    plt.xticks(years)
    plt.tight_layout()
    plt.savefig(output_path)
    plt.close()

def forecast(env_counts, category, num_years=3):
    existing_years = sorted(int(y) for y in env_counts.keys())
    years_str = [str(y) for y in existing_years]

    percentages = []
    valid_years = []

    for y in years_str:
        total = sum(env_counts[y].values())
        cat_val = env_counts[y].get(category, 0)
        percent = (cat_val / total) * 100 if total > 0 else 0
        if total > 0:
            percentages.append(percent)
            valid_years.append(int(y))

    if len(valid_years) == 0:
        print("Forecast skipped: No valid data for category:", category)
        return []

    X = np.array(valid_years).reshape(-1, 1)
    y = np.array(percentages)

    model = LinearRegression()
    model.fit(X, y)

    last_year = max(valid_years)
    future_years = [last_year + i for i in range(1, num_years + 1)]
    future_X = np.array(future_years).reshape(-1, 1)
    pred_percentages = model.predict(future_X)

    return list(zip(future_years, pred_percentages))

def generate_doc(env_counts, haze_counts, forecast_data, doc_path):
    doc = Document()
    doc.add_heading("Environmental Monitoring Report", 0)

    if os.path.exists("static/result_chart.png"):
        doc.add_picture("static/result_chart.png", width=Inches(6))
        doc.add_paragraph("")

    for year in sorted(env_counts, key=lambda x: int(x)):
        year = str(year)
        doc.add_heading(f"Year {year} Summary", level=1)
        total = sum(env_counts[year].values())
        for k, v in env_counts[year].items():
            percent = (v / total) * 100 if total > 0 else 0
            doc.add_paragraph(f"- {k}: {v} images ({percent:.2f}%)")

        doc.add_paragraph("\nHaze Summary:")
        haze_total = sum(haze_counts[year].values())
        for k, v in haze_counts[year].items():
            percent = (v / haze_total) * 100 if haze_total > 0 else 0
            doc.add_paragraph(f"- {k.capitalize()}: {v} images ({percent:.2f}%)")

    doc.add_heading("Forecasting (Green Cover %)", level=1)
    if forecast_data:
        for year, pred in forecast_data:
            doc.add_paragraph(f"- {year}: {round(pred, 2)}% expected green cover")
    else:
        doc.add_paragraph("Forecast not available due to insufficient data.")

    doc.add_heading("Policy Suggestions", level=1)
    all_env = sum((env_counts[y] for y in env_counts), Counter())
    all_haze = sum((haze_counts[y] for y in haze_counts), Counter())

    green_cover = all_env.get("Green Cover", 0)
    urban_area = all_env.get("Urban Area", 0)
    arid_zone = all_env.get("Arid Zone", 0)

    if urban_area > green_cover:
        doc.add_paragraph("Urbanization is increasing significantly.")
        doc.add_paragraph("- Promote urban afforestation (rooftop gardens, city trees).")
        doc.add_paragraph("- Enforce green buffer zones around cities.")
    else:
        doc.add_paragraph("Green Cover is maintained.")
        doc.add_paragraph("- Allow limited eco-friendly urban development.")
        doc.add_paragraph("- Focus on forest conservation programs.")

    if arid_zone > 0.2 * (green_cover + urban_area):
        doc.add_paragraph("Notable presence of Arid Zones.")
        doc.add_paragraph("- Launch soil restoration and irrigation initiatives.")

    hazy = all_haze.get('hazy', 0)
    clear = all_haze.get('clear', 0)

    if hazy > clear:
        doc.add_paragraph("Haze/Smog levels are high.")
        doc.add_paragraph("- Enforce emission regulations in industrial zones.")
        doc.add_paragraph("- Encourage public use of masks and cleaner transportation.")
    else:
        doc.add_paragraph("Air quality is generally good.")
        doc.add_paragraph("- Maintain pollution control policies and awareness campaigns.")

    doc.save(doc_path)
    print("DOCX report generated with updated policies and chart.")

def analyze_and_generate_report(upload_folder, result_folder, chart_path):
    env_data, haze_data = analyze_uploaded_images(upload_folder)
    generate_plot(env_data, haze_data, chart_path)
    forecast_data = forecast(env_data, category="Green Cover")
    report_path = os.path.join(result_folder, "final_report.docx")
    generate_doc(env_data, haze_data, forecast_data, report_path)
    print("Report and chart generated successfully.")
